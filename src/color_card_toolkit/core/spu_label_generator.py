from __future__ import annotations

from copy import deepcopy
from datetime import date, datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from openpyxl import load_workbook

from color_card_toolkit.core.resources import spu_label_template_path

LABEL_ROWS = 24
LABEL_COLUMNS = 6
LABELS_PER_PAGE = LABEL_ROWS * LABEL_COLUMNS


def read_spu_values_from_excel(
    excel_path: str | Path,
    *,
    start_row: int,
    start_column: int,
) -> list[str]:
    if start_row < 1 or start_column < 1:
        raise ValueError("起始行和起始列必须大于 0")

    workbook = load_workbook(filename=str(excel_path), read_only=True, data_only=True)
    try:
        sheet = workbook.worksheets[0]
        values: list[str] = []
        row = start_row
        while True:
            value = sheet.cell(row=row, column=start_column).value
            text = _format_cell_value(value)
            if not text:
                break
            values.append(text)
            row += 1
        return values
    finally:
        workbook.close()


def generate_spu_label_docx(
    values: Iterable[str],
    output_path: str | Path,
    template_path: str | Path | None = None,
) -> Path:
    labels = [str(value).strip() for value in values if str(value).strip()]
    if not labels:
        raise ValueError("没有可写入的不干胶内容")

    template = Path(template_path) if template_path else spu_label_template_path()
    if not template.exists():
        raise FileNotFoundError(f"找不到内置不干胶模板：{template}")

    document = Document(str(template))
    if not document.tables:
        raise ValueError("不干胶模板中未找到表格")

    template_table_xml = deepcopy(document.tables[0]._tbl)
    _keep_only_first_table(document)
    page_count = (len(labels) + LABELS_PER_PAGE - 1) // LABELS_PER_PAGE
    _ensure_table_count(document, template_table_xml, page_count)

    for page_index, table in enumerate(document.tables):
        start = page_index * LABELS_PER_PAGE
        _fill_label_table(table, labels[start : start + LABELS_PER_PAGE])
        if page_index > 0:
            _set_table_page_break_before(table)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    return output


def convert_spu_excel_to_docx(
    excel_path: str | Path,
    output_path: str | Path,
    *,
    start_row: int,
    start_column: int,
    template_path: str | Path | None = None,
) -> Path:
    values = read_spu_values_from_excel(
        excel_path,
        start_row=start_row,
        start_column=start_column,
    )
    return generate_spu_label_docx(values, output_path, template_path)


def _format_cell_value(value) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, bool):
        return "TRUE" if value else "FALSE"
    if isinstance(value, int):
        return str(value)
    if isinstance(value, float):
        return str(int(value)) if value.is_integer() else str(value)
    if isinstance(value, (datetime, date)):
        return value.strftime("%Y-%m-%d")
    return str(value).strip()


def _keep_only_first_table(document) -> None:
    body = document._body._element
    kept_table = False
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag == qn("w:tbl") and not kept_table:
            kept_table = True
            continue
        body.remove(child)


def _ensure_table_count(document, template_table_xml, count: int) -> None:
    body = document._body._element
    for _ in range(count - len(document.tables)):
        _insert_before_section_properties(body, deepcopy(template_table_xml))


def _insert_before_section_properties(body, element) -> None:
    section_properties = body.sectPr
    if section_properties is None:
        body.append(element)
    else:
        section_properties.addprevious(element)


def _fill_label_table(table: Table, labels: list[str]) -> None:
    for row_index in range(LABEL_ROWS):
        for column_index in range(LABEL_COLUMNS):
            cell = table.cell(row_index, column_index)
            index = row_index * LABEL_COLUMNS + column_index
            cell.text = labels[index] if index < len(labels) else ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_table_page_break_before(table: Table) -> None:
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph_properties = paragraph._p.get_or_add_pPr()
    if paragraph_properties.find(qn("w:pageBreakBefore")) is None:
        paragraph_properties.append(OxmlElement("w:pageBreakBefore"))
