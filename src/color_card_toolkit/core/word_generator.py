from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table

from color_card_toolkit.core.models import GroupedColorCard
from color_card_toolkit.core.resources import flat_template_path

CODE_ROW_INDICES = (1, 3, 5, 7)
PAGE_SIZE = 24


def generate_flat_template_docx(
    groups: list[GroupedColorCard],
    output_path: str | Path,
    template_path: str | Path | None = None,
) -> Path:
    if not groups:
        raise ValueError("没有可生成的色卡组")

    template = Path(template_path) if template_path else flat_template_path()
    if not template.exists():
        raise FileNotFoundError(f"找不到内置模板：{template}")

    pages = _build_pages(groups)
    document = Document(str(template))
    if not document.tables:
        raise ValueError("模板中未找到表格")

    template_table_xml = deepcopy(document.tables[0]._tbl)
    _keep_only_first_table(document)
    _ensure_table_count(document, template_table_xml, len(pages))

    for index, (table, (header, codes)) in enumerate(zip(document.tables, pages)):
        _fill_table(table, header, codes)
        if index > 0:
            _set_table_page_break_before(table)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    return output


def _build_pages(groups: list[GroupedColorCard]) -> list[tuple[str, list[str]]]:
    pages: list[tuple[str, list[str]]] = []
    for group in groups:
        for page_index, start in enumerate(range(0, len(group.color_codes), PAGE_SIZE), start=1):
            pages.append((f"{group.base_name} ({page_index})", group.color_codes[start : start + PAGE_SIZE]))
    return pages


def _ensure_table_count(document, template_table_xml, count: int) -> None:
    body = document._body._element
    for _ in range(count - len(document.tables)):
        _insert_before_section_properties(body, deepcopy(template_table_xml))


def _keep_only_first_table(document) -> None:
    body = document._body._element
    kept_table = False
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag == qn("w:tbl") and not kept_table:
            kept_table = True
            continue
        body.remove(child)


def _insert_before_section_properties(body, element) -> None:
    section_properties = body.sectPr
    if section_properties is None:
        body.append(element)
    else:
        section_properties.addprevious(element)


def _fill_table(table: Table, header: str, codes: list[str]) -> None:
    table.cell(0, 0).text = header
    _align_cell(table.cell(0, 0), bold=True)

    for row_index in CODE_ROW_INDICES:
        for column_index in range(6):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            _align_cell(cell, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    for index, code in enumerate(codes[:PAGE_SIZE]):
        row_index = CODE_ROW_INDICES[index // 6]
        column_index = index % 6
        cell = table.cell(row_index, column_index)
        cell.text = str(code)
        _align_cell(cell, alignment=WD_ALIGN_PARAGRAPH.LEFT)


def _align_cell(cell, *, bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        for run in paragraph.runs:
            run.bold = bold


def _set_table_page_break_before(table: Table) -> None:
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph_properties = paragraph._p.get_or_add_pPr()
    if paragraph_properties.find(qn("w:pageBreakBefore")) is None:
        paragraph_properties.append(OxmlElement("w:pageBreakBefore"))
