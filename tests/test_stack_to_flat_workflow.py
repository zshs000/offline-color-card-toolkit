from __future__ import annotations

from pathlib import Path

from docx import Document

from color_card_toolkit.core.grouping import group_recognition_results
from color_card_toolkit.core.models import GroupedColorCard, OcrBlock
from color_card_toolkit.core.resources import flat_template_path
from color_card_toolkit.core.word_generator import generate_flat_template_docx


def test_stack_to_flat_service_workflow_generates_paginated_word(tmp_path: Path) -> None:
    codes = [str(number).zfill(2) for number in range(1, 27)]
    groups = [GroupedColorCard("PU88", codes)]
    output_path = tmp_path / "result.docx"

    generate_flat_template_docx(groups, output_path, flat_template_path())

    document = Document(str(output_path))
    assert len(document.tables) == 2
    assert document.tables[0].cell(0, 0).text == "PU88 (1)"
    assert document.tables[1].cell(0, 0).text == "PU88 (2)"
    assert document.tables[1].cell(1, 0).text == "25"
    assert document.tables[1].cell(1, 1).text == "26"

