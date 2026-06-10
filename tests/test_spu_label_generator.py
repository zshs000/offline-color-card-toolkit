from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook

from color_card_toolkit.core.resources import spu_label_template_path
from color_card_toolkit.core.spu_label_generator import (
    generate_spu_label_docx,
    read_spu_values_from_excel,
)


def test_read_spu_values_from_first_sheet_selected_column_until_blank(tmp_path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "first"
    sheet["A2"] = "不要读取"
    sheet["B2"] = "SPU-001"
    sheet["B3"] = "SPU-002"
    sheet["B4"] = "SPU-003"
    sheet["B5"] = None
    sheet["B6"] = "不要读取"
    second = workbook.create_sheet("second")
    second["B2"] = "不要读取第二页"
    excel_path = tmp_path / "sample.xlsx"
    workbook.save(excel_path)

    assert read_spu_values_from_excel(excel_path, start_row=2, start_column=2) == [
        "SPU-001",
        "SPU-002",
        "SPU-003",
    ]


def test_generate_spu_label_docx_duplicates_first_template_table_for_extra_pages(tmp_path: Path) -> None:
    output_path = tmp_path / "spu.docx"
    values = [f"SPU-{index:03d}" for index in range(1, 146)]

    generate_spu_label_docx(values, output_path, spu_label_template_path())

    document = Document(str(output_path))
    assert len(document.tables) == 2
    assert len(document.tables[0].rows) == 24
    assert len(document.tables[0].columns) == 6
    assert document.tables[0].cell(0, 0).text == "SPU-001"
    assert document.tables[0].cell(0, 5).text == "SPU-006"
    assert document.tables[0].cell(23, 5).text == "SPU-144"
    assert document.tables[1].cell(0, 0).text == "SPU-145"
