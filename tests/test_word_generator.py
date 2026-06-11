from __future__ import annotations

from pathlib import Path
from xml.etree import ElementTree as ET
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from color_card_toolkit.core.models import GroupedColorCard
from color_card_toolkit.core.word_generator import CODE_ROW_INDICES, generate_flat_template_docx

WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def test_word_generator_fills_header_and_24_code_slots(tmp_path: Path) -> None:
    template = Path("docs/转平贴底纸模板.docx")
    output = tmp_path / "output.docx"
    codes = [str(number).zfill(2) for number in range(1, 27)]

    generate_flat_template_docx([GroupedColorCard("PU88", codes)], output, template)

    document = Document(str(output))
    assert len(document.tables) == 2
    assert document.tables[0].cell(0, 0).text == "PU88 (1)"
    assert document.tables[1].cell(0, 0).text == "PU88 (2)"
    assert document.tables[0].cell(0, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT

    first_page_codes = []
    for row_index in CODE_ROW_INDICES:
        first_page_codes.extend(document.tables[0].cell(row_index, column).text for column in range(6))
    assert first_page_codes == codes[:24]
    assert document.tables[0].cell(1, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert document.tables[0].cell(7, 5).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT

    second_page_codes = []
    for row_index in CODE_ROW_INDICES:
        second_page_codes.extend(document.tables[1].cell(row_index, column).text for column in range(6))
    assert second_page_codes[:2] == ["25", "26"]
    assert all(value == "" for value in second_page_codes[2:])


def test_word_generator_uses_table_page_breaks_without_blank_break_paragraphs(tmp_path: Path) -> None:
    template = Path("docs/转平贴底纸模板.docx")
    output = tmp_path / "output.docx"

    generate_flat_template_docx([GroupedColorCard("PU88", ["01", "02"] * 13)], output, template)

    body_children = _document_body_children(output)
    child_tags = [_local_name(child.tag) for child in body_children]
    standalone_page_breaks = [
        child
        for child in body_children
        if _local_name(child.tag) == "p" and child.findall(".//w:br[@w:type='page']", WORD_NS)
    ]
    page_break_before_count = sum(
        len(table.findall(".//w:pageBreakBefore", WORD_NS))
        for table in body_children
        if _local_name(table.tag) == "tbl"
    )

    assert child_tags == ["tbl", "tbl", "sectPr"]
    assert standalone_page_breaks == []
    assert page_break_before_count == 1


def test_word_generator_ignores_extra_template_pages(tmp_path: Path) -> None:
    template = tmp_path / "two_page_template.docx"
    source = Document()
    source.add_table(rows=9, cols=6)
    source.add_page_break()
    source.add_table(rows=9, cols=6).cell(0, 0).text = "EXTRA TEMPLATE PAGE"
    source.save(str(template))

    output = tmp_path / "output.docx"

    generate_flat_template_docx([GroupedColorCard("PU88", ["01"])], output, template)

    document = Document(str(output))
    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "PU88 (1)"
    assert "EXTRA TEMPLATE PAGE" not in "\n".join(table.cell(0, 0).text for table in document.tables)


def _document_body_children(path: Path) -> list[ET.Element]:
    with ZipFile(path) as package:
        document_xml = package.read("word/document.xml")
    root = ET.fromstring(document_xml)
    body = root.find("w:body", WORD_NS)
    assert body is not None
    return list(body)


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]
